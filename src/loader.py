from pypdf import PdfReader
from pathlib import Path
import re
from typing import List, Optional

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class DocumentLoader:
    """Load text from multiple document types. Keep functions small and testable."""

    @staticmethod
    def load_pdf(path: str) -> str:
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {path}")
        reader = PdfReader(path)
        pages = []
        for page in reader.pages:
            # extract_text() may return None
            txt = page.extract_text()
            if txt:
                pages.append(txt)
        text = "\n".join(pages).strip()
        if not text:
            raise ValueError("No extractable text found in PDF")
        return DocumentLoader._preprocess_text(text)

    @staticmethod
    def load_txt(path: str) -> str:
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {path}")
        text = p.read_text(encoding="utf-8")
        return DocumentLoader._preprocess_text(text)

    @staticmethod
    def load_docx(path: str) -> str:
        """Load text from Word document (.docx)"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {path}")
        
        try:
            doc = Document(path)
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Add table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            text = "\n".join(paragraphs)
            if not text.strip():
                raise ValueError("No extractable text found in DOCX")
            
            return DocumentLoader._preprocess_text(text)
        except Exception as e:
            raise ValueError(f"Error reading DOCX file: {str(e)}")

    @staticmethod
    def load_md(path: str) -> str:
        """Load text from Markdown file (.md)"""
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {path}")
        
        text = p.read_text(encoding="utf-8")
        # Remove markdown formatting for cleaner text processing
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)  # Remove header marks
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)  # Remove italic  
        text = re.sub(r'`(.*?)`', r'\1', text)  # Remove inline code
        text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)  # Remove code blocks
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Convert links to text
        text = re.sub(r'^\s*[\*\-\+]\s+', '', text, flags=re.MULTILINE)  # Remove list markers
        
        return DocumentLoader._preprocess_text(text)

    @staticmethod
    def load_document(path: str) -> str:
        """Auto-detect document type and load appropriately"""
        p = Path(path)
        suffix = p.suffix.lower()
        
        if suffix == '.pdf':
            return DocumentLoader.load_pdf(path)
        elif suffix == '.txt':
            return DocumentLoader.load_txt(path)
        elif suffix == '.docx':
            return DocumentLoader.load_docx(path)
        elif suffix in ['.md', '.markdown']:
            return DocumentLoader.load_md(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}. Supported: .pdf, .txt, .docx, .md")

    @staticmethod
    def _preprocess_text(text: str) -> str:
        """Clean and preprocess text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might cause issues but keep important punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}\"\'\/\@\#\$\%\&\*\+\=\n]', ' ', text)
        
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        
        # Strip and ensure single spaces
        text = ' '.join(text.split())
        
        return text.strip()

    @staticmethod
    def get_supported_extensions() -> List[str]:
        """Get list of supported file extensions"""
        extensions = ['.pdf', '.txt', '.md', '.markdown']
        if DOCX_AVAILABLE:
            extensions.append('.docx')
        return extensions

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 3000, overlap: int = 200) -> list:
        """Rudimentary chunker by characters (safe for tokens approx). Returns list of chunks."""
        if chunk_size <= overlap:
            raise ValueError("chunk_size must be greater than overlap")
        chunks = []
        start = 0
        length = len(text)
        while start < length:
            end = start + chunk_size
            chunks.append(text[start:end])
            start = end - overlap
        return chunks

    @staticmethod
    def chunk_text_smart(text: str, chunk_size: int = 3000, overlap: int = 200) -> List[str]:
        """
        Sentence-aware chunking that tries to break at sentence boundaries
        """
        if chunk_size <= overlap:
            raise ValueError("chunk_size must be greater than overlap")
        
        # Split into sentences (basic approach)
        sentences = DocumentLoader._split_into_sentences(text)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            # If adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap from previous chunk
                if overlap > 0 and chunks:
                    overlap_text = DocumentLoader._get_overlap_text(current_chunk, overlap)
                    current_chunk = overlap_text + " " + sentence
                else:
                    current_chunk = sentence
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        # Add the last chunk if it has content
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks

    @staticmethod
    def _split_into_sentences(text: str) -> List[str]:
        """Split text into sentences using regex"""
        # Simple sentence splitting - can be improved with nltk/spacy
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]

    @staticmethod
    def _get_overlap_text(text: str, overlap_chars: int) -> str:
        """Get the last N characters for overlap, trying to break at word boundaries"""
        if len(text) <= overlap_chars:
            return text
        
        overlap_text = text[-overlap_chars:]
        
        # Try to start at a word boundary
        space_idx = overlap_text.find(' ')
        if space_idx != -1:
            return overlap_text[space_idx:].strip()
        
        return overlap_text
